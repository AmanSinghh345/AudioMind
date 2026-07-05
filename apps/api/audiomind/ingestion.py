from __future__ import annotations

import hashlib
import csv
import io
import re
import shutil
import uuid
from pathlib import Path

import fitz
import pytesseract
from docx import Document
from PIL import Image

from .chunking import PageAwareChunker
from .config import Settings
from .embeddings import EmbeddingService
from .models import ExtractedPage
from .repository import Repository
from .vector_store import ChromaVectorStore


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".csv", ".png", ".jpg", ".jpeg"}


class IngestionService:
    def __init__(
        self, settings: Settings, repository: Repository,
        embeddings: EmbeddingService, vector_store: ChromaVectorStore,
    ):
        self.settings = settings
        self.repository = repository
        self.embeddings = embeddings
        self.vector_store = vector_store
        self.chunker = PageAwareChunker(settings.chunk_size, settings.chunk_overlap)
        if settings.tesseract_cmd and Path(settings.tesseract_cmd).exists():
            pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd

    def ingest_bytes(
        self, collection_id: str, filename: str, content: bytes, use_ocr: bool = True
    ) -> dict:
        safe_name = Path(filename).name
        extension = Path(safe_name).suffix.lower()
        self._validate(safe_name, extension, content)
        content_hash = hashlib.sha256(content).hexdigest()
        existing = self.repository.find_document(collection_id, safe_name)
        if existing and existing["content_hash"] == content_hash and existing["status"] == "ready":
            return {**existing, "deduplicated": True}

        if existing:
            self.vector_store.delete_document(existing["id"])
            old_path = Path(existing["stored_path"])
            self.repository.delete_document(existing["id"])
            if old_path.exists():
                old_path.unlink()

        document_id = str(uuid.uuid4())
        document_dir = self.settings.upload_path / collection_id
        document_dir.mkdir(parents=True, exist_ok=True)
        stored_path = document_dir / f"{document_id}{extension}"
        stored_path.write_bytes(content)
        self.repository.create_document(
            document_id, collection_id, safe_name, extension, content_hash, str(stored_path)
        )

        try:
            pages = self.extract_pages(stored_path, use_ocr=use_ocr)
            if not pages or not any(page.text.strip() for page in pages):
                raise ValueError("No readable text was found in the document")
            chunks = self.chunker.chunk(pages, document_id, collection_id, safe_name)
            vectors = self.embeddings.embed([chunk.text for chunk in chunks])
            self.vector_store.upsert(chunks, vectors)
            self.repository.finish_document(document_id, len(pages), chunks)
            return {
                "id": document_id,
                "collection_id": collection_id,
                "filename": safe_name,
                "status": "ready",
                "page_count": len(pages),
                "chunk_count": len(chunks),
                "embedding_method": self.embeddings.method,
                "deduplicated": False,
            }
        except Exception as exc:
            self.repository.fail_document(document_id, str(exc))
            self.vector_store.delete_document(document_id)
            raise

    def delete(self, document_id: str) -> None:
        document = self.repository.get_document(document_id)
        if not document:
            return
        self.vector_store.delete_document(document_id)
        self.repository.delete_document(document_id)
        path = Path(document["stored_path"])
        if path.exists():
            path.unlink()

    def extract_pages(self, path: Path, use_ocr: bool = True) -> list[ExtractedPage]:
        extension = path.suffix.lower()
        if extension == ".pdf":
            return self._pdf_pages(path, use_ocr)
        if extension == ".docx":
            return self._docx_pages(path)
        if extension in {".txt", ".md"}:
            return [ExtractedPage(1, path.read_text(encoding="utf-8", errors="replace"))]
        if extension == ".csv":
            raw = path.read_text(encoding="utf-8-sig", errors="replace")
            rows = list(csv.reader(io.StringIO(raw)))
            text = "\n".join(" | ".join(cell.strip() for cell in row) for row in rows)
            return [ExtractedPage(1, text, "Structured data")]
        if extension in {".png", ".jpg", ".jpeg"}:
            text = pytesseract.image_to_string(Image.open(path)).strip()
            return [ExtractedPage(1, text, "Scanned notes")]
        raise ValueError(f"Unsupported file type: {extension}")

    def _pdf_pages(self, path: Path, use_ocr: bool) -> list[ExtractedPage]:
        pages: list[ExtractedPage] = []
        with fitz.open(path) as document:
            for index, page in enumerate(document):
                text = page.get_text("text", sort=True).strip()
                if use_ocr and len(text) < 40:
                    pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    image = Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
                    ocr_text = pytesseract.image_to_string(image).strip()
                    if len(ocr_text) > len(text):
                        text = ocr_text
                text = self._mark_likely_headings(text)
                pages.append(ExtractedPage(index + 1, text))
        return pages

    @staticmethod
    def _docx_pages(path: Path) -> list[ExtractedPage]:
        document = Document(path)
        lines: list[str] = []
        chapter = "Document"
        for paragraph in document.paragraphs:
            value = paragraph.text.strip()
            if not value:
                continue
            if paragraph.style and paragraph.style.name.startswith("Heading"):
                chapter = value
                lines.append(f"## {value}")
            else:
                lines.append(value)
        for table in document.tables:
            for row in table.rows:
                lines.append(" | ".join(cell.text.strip() for cell in row.cells))
        return [ExtractedPage(1, "\n\n".join(lines), chapter)]

    @staticmethod
    def _mark_likely_headings(text: str) -> str:
        lines = []
        for line in text.splitlines():
            stripped = line.strip()
            if stripped and len(stripped) < 80 and not re.search(r"[.!?]$", stripped):
                if stripped.isupper() or re.match(r"^(chapter|unit|section)\s+\w+", stripped, re.I):
                    stripped = f"## {stripped}"
            lines.append(stripped)
        return "\n".join(lines)

    def _validate(self, filename: str, extension: str, content: bytes) -> None:
        if not filename or extension not in ALLOWED_EXTENSIONS:
            raise ValueError(f"Unsupported file type. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}")
        if not content:
            raise ValueError("The uploaded file is empty")
        if len(content) > self.settings.max_upload_mb * 1024 * 1024:
            raise ValueError(f"File exceeds the {self.settings.max_upload_mb} MB upload limit")
